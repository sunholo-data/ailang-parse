"""Generate challenge benchmark files that test DocParse's known gaps.

These files exercise OOXML features we DON'T handle well yet,
giving us an honest benchmark baseline to improve against.

Usage: uv run benchmarks/create_challenge_files.py
"""

import os
from pathlib import Path

OUTPUT_DIR = Path("data/test_files/challenge")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_challenge_formatting():
    """DOCX with bold/italic carrying semantic meaning.

    Tests: w:rPr (§17.3.2) — bold, italic, underline, strikethrough.
    Expected gap: All formatting stripped, semantic meaning lost.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_UNDERLINE

    doc = Document()
    doc.add_heading("Scientific Paper: Gene Expression Analysis", level=1)

    # Paragraph with semantically meaningful formatting
    p = doc.add_paragraph()
    p.add_run("The gene ")
    run_gene = p.add_run("BRCA1")
    run_gene.bold = True
    run_gene.italic = True
    p.add_run(" (also known as ")
    run_alt = p.add_run("breast cancer type 1")
    run_alt.italic = True
    p.add_run(") is located on chromosome 17. The protein product, ")
    run_protein = p.add_run("BRCA1")
    run_protein.bold = True
    p.add_run(", acts as a tumor suppressor.")

    # Book references with italic titles
    p2 = doc.add_paragraph()
    p2.add_run("As discussed in ")
    run_book = p2.add_run("The Emperor of All Maladies")
    run_book.italic = True
    p2.add_run(" by Siddhartha Mukherjee, and ")
    run_book2 = p2.add_run("The Gene: An Intimate History")
    run_book2.italic = True
    p2.add_run(", the understanding of cancer genetics has evolved dramatically.")

    # Technical terms with specific formatting
    p3 = doc.add_paragraph()
    p3.add_run("The ")
    run_term = p3.add_run("p53")
    run_term.bold = True
    p3.add_run(" pathway regulates ")
    run_term2 = p3.add_run("apoptosis")
    run_term2.underline = True
    p3.add_run(". Mutations in ")
    run_deprecated = p3.add_run("TP53")
    run_deprecated.font.strike = True
    p3.add_run(" (now called ")
    run_new = p3.add_run("TP53")
    run_new.bold = True
    p3.add_run(") are found in ~50% of cancers.")

    # Color-coded status indicators
    p4 = doc.add_paragraph()
    run_pass = p4.add_run("PASSED")
    run_pass.font.color.rgb = RGBColor(0, 128, 0)
    run_pass.bold = True
    p4.add_run(" — Sample A; ")
    run_fail = p4.add_run("FAILED")
    run_fail.font.color.rgb = RGBColor(255, 0, 0)
    run_fail.bold = True
    p4.add_run(" — Sample B; ")
    run_warn = p4.add_run("WARNING")
    run_warn.font.color.rgb = RGBColor(255, 165, 0)
    run_warn.bold = True
    p4.add_run(" — Sample C")

    path = OUTPUT_DIR / "challenge_formatting.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def create_challenge_numbering():
    """DOCX with multi-level numbered outline.

    Tests: word/numbering.xml (§17.9) — abstract numbering definitions.
    Expected gap: List levels wrong, ordered/unordered guessed incorrectly.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_heading("Project Plan: Multi-Level Outline", level=1)

    # Create proper multi-level list using List Number styles
    items = [
        (0, "Introduction", "List Number"),
        (1, "Background", "List Number 2"),
        (1, "Objectives", "List Number 2"),
        (2, "Primary objective", "List Number 3"),
        (2, "Secondary objective", "List Number 3"),
        (0, "Methodology", "List Number"),
        (1, "Data collection", "List Number 2"),
        (2, "Survey design", "List Number 3"),
        (2, "Sample size calculation", "List Number 3"),
        (2, "Recruitment strategy", "List Number 3"),
        (1, "Analysis approach", "List Number 2"),
        (0, "Results", "List Number"),
        (0, "Discussion", "List Number"),
        (1, "Limitations", "List Number 2"),
        (1, "Future work", "List Number 2"),
    ]

    for level, text, style_name in items:
        p = doc.add_paragraph(text, style=style_name)

    # Also add bullet lists at multiple levels
    doc.add_heading("Requirements (Bullets)", level=2)
    bullet_items = [
        (0, "Must have features", "List Bullet"),
        (1, "User authentication", "List Bullet 2"),
        (1, "Data export", "List Bullet 2"),
        (2, "CSV format", "List Bullet 3"),
        (2, "JSON format", "List Bullet 3"),
        (0, "Nice to have features", "List Bullet"),
        (1, "Dark mode", "List Bullet 2"),
        (1, "Keyboard shortcuts", "List Bullet 2"),
    ]

    for level, text, style_name in bullet_items:
        p = doc.add_paragraph(text, style=style_name)

    path = OUTPUT_DIR / "challenge_numbering.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def create_challenge_styles():
    """DOCX with custom heading styles.

    Tests: word/styles.xml (§17.7) — style definitions and inheritance.
    Expected gap: Custom styles detected as "Normal" text, not headings.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Add custom styles that should be recognized as headings
    styles = doc.styles

    # Create "Chapter Title" style based on Heading 1
    chapter_style = styles.add_style("ChapterTitle", WD_STYLE_TYPE.PARAGRAPH)
    chapter_style.base_style = styles["Heading 1"]
    chapter_style.font.size = Pt(28)
    chapter_style.font.bold = True

    # Create "Section Header" style based on Heading 2
    section_style = styles.add_style("SectionHeader", WD_STYLE_TYPE.PARAGRAPH)
    section_style.base_style = styles["Heading 2"]
    section_style.font.size = Pt(18)

    # Create "Subsection" style based on Heading 3
    sub_style = styles.add_style("Subsection", WD_STYLE_TYPE.PARAGRAPH)
    sub_style.base_style = styles["Heading 3"]
    sub_style.font.size = Pt(14)

    # Use the custom styles
    doc.add_paragraph("Annual Report 2025", style="ChapterTitle")
    doc.add_paragraph("This document uses custom heading styles that inherit from built-in headings.")

    doc.add_paragraph("Financial Overview", style="SectionHeader")
    doc.add_paragraph("Revenue grew 15% year-over-year driven by cloud services.")

    doc.add_paragraph("Revenue Breakdown", style="Subsection")
    doc.add_paragraph("Cloud: $45M, On-prem: $12M, Services: $8M")

    doc.add_paragraph("Operational Highlights", style="SectionHeader")
    doc.add_paragraph("Headcount increased by 200 across engineering and sales.")

    doc.add_paragraph("Engineering", style="Subsection")
    doc.add_paragraph("Shipped 3 major releases with 99.9% uptime.")

    # Also use Title and Subtitle (built-in but not "Heading1"-"Heading6")
    doc.add_paragraph("Appendix: Supplementary Data", style="Title")
    doc.add_paragraph("Detailed financial tables", style="Subtitle")

    path = OUTPUT_DIR / "challenge_styles.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def create_challenge_fields():
    """DOCX with field codes: TOC, cross-references, page numbers.

    Tests: w:fldSimple, w:fldChar (§17.16) — field codes and display text.
    Expected gap: Fields completely invisible, TOC entries lost.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_heading("Document with Field Codes", level=1)

    # Add a paragraph with a simple HYPERLINK field
    p = doc.add_paragraph()
    # Create w:hyperlink element with r:id
    hyperlink = OxmlElement("w:hyperlink")
    # We can't easily set r:id without rels, so use w:fldSimple instead

    # Add a fldSimple for a date field
    p2 = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " DATE \\@ \"MMMM d, yyyy\" ")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "March 28, 2026"
    run.append(text)
    fld.append(run)
    p2._element.append(fld)

    # Add content with heading references
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This section introduces the topic.")

    doc.add_heading("Methods", level=2)
    doc.add_paragraph("See the introduction above for context.")

    # Add a NUMPAGES field
    p3 = doc.add_paragraph("Total pages in document: ")
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), " NUMPAGES ")
    run2 = OxmlElement("w:r")
    text2 = OxmlElement("w:t")
    text2.text = "3"
    run2.append(text2)
    fld2.append(run2)
    p3._element.append(fld2)

    # Add a FILENAME field
    p4 = doc.add_paragraph("File: ")
    fld3 = OxmlElement("w:fldSimple")
    fld3.set(qn("w:instr"), " FILENAME ")
    run3 = OxmlElement("w:r")
    text3 = OxmlElement("w:t")
    text3.text = "challenge_fields.docx"
    run3.append(text3)
    fld3.append(run3)
    p4._element.append(fld3)

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("This document demonstrates field code extraction.")

    path = OUTPUT_DIR / "challenge_fields.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def create_challenge_hyperlinks():
    """DOCX with hyperlinks that have URL targets.

    Tests: w:hyperlink + r:id → relationship resolution (§17.16.22).
    Expected gap: Link text extracted but URLs discarded.
    """
    from docx import Document
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    import docx.opc.constants

    doc = Document()
    doc.add_heading("Document with Hyperlinks", level=1)

    # python-docx doesn't have built-in hyperlink support, use add_hyperlink helper
    def add_hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._element.append(hyperlink)

    p1 = doc.add_paragraph("Visit our website: ")
    add_hyperlink(p1, "https://www.sunholo.com", "Sunholo Homepage")

    p2 = doc.add_paragraph("Documentation is available at ")
    add_hyperlink(p2, "https://www.sunholo.com/docparse", "DocParse Docs")
    p2.add_run(".")

    p3 = doc.add_paragraph("For issues, see ")
    add_hyperlink(p3, "https://github.com/sunholo-data/docparse/issues", "GitHub Issues")
    p3.add_run(" or email ")
    add_hyperlink(p3, "mailto:support@sunholo.com", "support@sunholo.com")
    p3.add_run(".")

    doc.add_heading("References", level=2)
    refs = [
        ("https://ecma-international.org/publications-and-standards/standards/ecma-376/", "ECMA-376 Standard"),
        ("https://docs.oasis-open.org/office/v1.2/", "ODF 1.2 Specification"),
        ("https://www.w3.org/TR/html52/", "HTML 5.2 Specification"),
    ]
    for url, text in refs:
        p = doc.add_paragraph()
        add_hyperlink(p, url, text)

    path = OUTPUT_DIR / "challenge_hyperlinks.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def create_challenge_merged_cells_xlsx():
    """XLSX with merged cell regions.

    Tests: <mergeCells> element (§18.3.1.55).
    Expected gap: Merges not detected, data appears in wrong cells.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Sales Report"

    # Header row with merged cells
    ws.merge_cells("A1:D1")
    ws["A1"] = "Q1 2026 Sales Report"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    # Sub-headers with merges
    ws.merge_cells("A3:A4")
    ws["A3"] = "Region"
    ws.merge_cells("B3:C3")
    ws["B3"] = "Revenue"
    ws["B4"] = "Actual"
    ws["C4"] = "Target"
    ws["D3"] = "Status"
    ws.merge_cells("D3:D4")

    # Data rows
    data = [
        ("North America", 45000, 40000, "Above Target"),
        ("Europe", 32000, 35000, "Below Target"),
        ("Asia Pacific", 28000, 25000, "Above Target"),
        ("Latin America", 12000, 15000, "Below Target"),
    ]
    for i, (region, actual, target, status) in enumerate(data, start=5):
        ws.cell(row=i, column=1, value=region)
        ws.cell(row=i, column=2, value=actual)
        ws.cell(row=i, column=3, value=target)
        ws.cell(row=i, column=4, value=status)

    # Total row with merge
    ws.merge_cells("A9:A9")
    ws["A9"] = "Total"
    ws["A9"].font = Font(bold=True)
    ws["B9"] = 117000
    ws["C9"] = 115000

    # Second sheet with more complex merges
    ws2 = wb.create_sheet("Matrix")
    # 3x3 merge in top-left
    ws2.merge_cells("A1:C3")
    ws2["A1"] = "Merged 3x3 Block"
    # 2x1 merge
    ws2.merge_cells("D1:E1")
    ws2["D1"] = "Merged 2 cols"
    # 1x2 merge
    ws2.merge_cells("D2:D3")
    ws2["D2"] = "Merged 2 rows"
    ws2["E2"] = "Normal"
    ws2["E3"] = "Normal"

    path = OUTPUT_DIR / "challenge_merged_cells.xlsx"
    wb.save(str(path))
    print(f"  Created {path}")


def create_challenge_real_world():
    """DOCX that mixes many features like a real-world report.

    Tests: Combined extraction of headings, tables, lists, formatting,
    hyperlinks, and structural elements in a realistic document.
    Expected gap: Multiple partial failures compound.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants

    doc = Document()

    # Title page content
    doc.add_paragraph("Quarterly Performance Review", style="Title")
    doc.add_paragraph("Engineering Division — Q1 2026", style="Subtitle")
    doc.add_paragraph("")

    # TOC-like field
    p_toc = doc.add_paragraph("Table of Contents")
    p_toc.style = doc.styles["Heading 1"]

    # Simulated TOC entries (in real docs these are field codes)
    toc_entries = [
        "1. Executive Summary .......................... 2",
        "2. Key Metrics ................................ 3",
        "3. Team Performance ........................... 5",
        "4. Risks and Mitigations ...................... 7",
    ]
    for entry in toc_entries:
        doc.add_paragraph(entry)

    # Executive Summary with formatted text
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Q1 2026 was a ")
    strong = p.add_run("strong quarter")
    strong.bold = True
    p.add_run(" for the engineering division. Key highlights include:")

    # Bullet list
    doc.add_paragraph("Shipped DocParse v0.9.0 with agent-friendly API", style="List Bullet")
    doc.add_paragraph("Achieved 99.95% API uptime (target: 99.9%)", style="List Bullet")
    doc.add_paragraph("Reduced p50 latency from 45ms to 11ms", style="List Bullet")

    # Key Metrics table
    doc.add_heading("Key Metrics", level=1)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Metric", "Q4 2025", "Q1 2026", "Change"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    metrics = [
        ("API Requests/day", "12,400", "28,700", "+131%"),
        ("Avg Latency (ms)", "45", "11", "-76%"),
        ("Error Rate", "0.12%", "0.05%", "-58%"),
        ("Active Users", "340", "892", "+162%"),
    ]
    for row_idx, (metric, q4, q1, change) in enumerate(metrics, start=1):
        table.rows[row_idx].cells[0].text = metric
        table.rows[row_idx].cells[1].text = q4
        table.rows[row_idx].cells[2].text = q1
        table.rows[row_idx].cells[3].text = change

    # Team section with sub-headings
    doc.add_heading("Team Performance", level=1)
    doc.add_heading("Backend Team", level=2)
    p2 = doc.add_paragraph()
    p2.add_run("Led by ")
    lead = p2.add_run("Sarah Chen")
    lead.bold = True
    p2.add_run(". The backend team delivered the ")
    feature = p2.add_run("capability manifest")
    feature.italic = True
    p2.add_run(" and ")
    feature2 = p2.add_run("request replay")
    feature2.italic = True
    p2.add_run(" features on schedule.")

    # Numbered list
    doc.add_paragraph("API key management system", style="List Number")
    doc.add_paragraph("Device authentication flow (RFC 8628)", style="List Number")
    doc.add_paragraph("Billing entitlement enforcement", style="List Number")

    doc.add_heading("Frontend Team", level=2)
    doc.add_paragraph("Launched the interactive API playground with Swagger UI integration.")

    # Risks section
    doc.add_heading("Risks and Mitigations", level=1)
    risk_table = doc.add_table(rows=4, cols=3)
    risk_table.style = "Table Grid"
    risk_headers = ["Risk", "Severity", "Mitigation"]
    for i, h in enumerate(risk_headers):
        risk_table.rows[0].cells[i].text = h
    risks = [
        ("AILANG codegen bugs block Go compilation", "High", "Track 3 blocking bugs, workaround with interpreter"),
        ("Moby Dick OOM on large documents", "Medium", "Implement streaming parser with FS budget increase"),
        ("Competitor launches Office parsing", "Low", "Maintain benchmark lead, focus on unique features"),
    ]
    for row_idx, (risk, severity, mitigation) in enumerate(risks, start=1):
        risk_table.rows[row_idx].cells[0].text = risk
        risk_table.rows[row_idx].cells[1].text = severity
        risk_table.rows[row_idx].cells[2].text = mitigation

    # Add a hyperlink
    def add_hyperlink(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._element.append(hyperlink)

    doc.add_heading("References", level=1)
    p_ref = doc.add_paragraph("Full API documentation: ")
    add_hyperlink(p_ref, "https://www.sunholo.com/docparse/api", "DocParse API Reference")

    path = OUTPUT_DIR / "challenge_real_world.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def create_challenge_equations():
    """DOCX with Office Math equations (OMML).

    Tests: m:oMath, m:oMathPara (§22.1).
    Expected gap: Equations produce empty/corrupted text.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_heading("Mathematical Notation", level=1)

    doc.add_paragraph("The following equations demonstrate Office Math (OMML) content:")

    # Create an inline equation: E = mc²
    p = doc.add_paragraph("Einstein's mass-energy equivalence: ")
    omath = OxmlElement("m:oMath")
    # E
    r1 = OxmlElement("m:r")
    t1 = OxmlElement("m:t")
    t1.text = "E"
    r1.append(t1)
    omath.append(r1)
    # =
    r2 = OxmlElement("m:r")
    t2 = OxmlElement("m:t")
    t2.text = "="
    r2.append(t2)
    omath.append(r2)
    # mc²
    r3 = OxmlElement("m:r")
    t3 = OxmlElement("m:t")
    t3.text = "m"
    r3.append(t3)
    omath.append(r3)
    # Superscript for c²
    sSup = OxmlElement("m:sSup")
    e_elem = OxmlElement("m:e")
    r4 = OxmlElement("m:r")
    t4 = OxmlElement("m:t")
    t4.text = "c"
    r4.append(t4)
    e_elem.append(r4)
    sSup.append(e_elem)
    sup_elem = OxmlElement("m:sup")
    r5 = OxmlElement("m:r")
    t5 = OxmlElement("m:t")
    t5.text = "2"
    r5.append(t5)
    sup_elem.append(r5)
    sSup.append(sup_elem)
    omath.append(sSup)
    p._element.append(omath)

    # Create a display equation: quadratic formula
    p2 = doc.add_paragraph("The quadratic formula:")
    oMathPara = OxmlElement("m:oMathPara")
    omath2 = OxmlElement("m:oMath")
    # x = (-b ± √(b²-4ac)) / 2a — simplified as text in m:r elements
    for text_part in ["x", "=", "(-b±√(b²-4ac))", "/", "2a"]:
        r = OxmlElement("m:r")
        t = OxmlElement("m:t")
        t.text = text_part
        r.append(t)
        omath2.append(r)
    oMathPara.append(omath2)
    p2._element.append(oMathPara)

    doc.add_paragraph("A well-formed parser should extract at least the text content of these equations.")

    path = OUTPUT_DIR / "challenge_equations.docx"
    doc.save(str(path))
    print(f"  Created {path}")


def main():
    print("Creating challenge benchmark files...")
    print(f"Output directory: {OUTPUT_DIR}")
    print()

    generators = [
        ("Formatting (bold/italic semantics)", create_challenge_formatting),
        ("Numbering (multi-level lists)", create_challenge_numbering),
        ("Custom styles", create_challenge_styles),
        ("Field codes (TOC, dates, page numbers)", create_challenge_fields),
        ("Hyperlinks with URLs", create_challenge_hyperlinks),
        ("XLSX merged cells", create_challenge_merged_cells_xlsx),
        ("Equations (OMML)", create_challenge_equations),
        ("Real-world mixed document", create_challenge_real_world),
    ]

    for name, gen_func in generators:
        print(f"  [{name}]")
        try:
            gen_func()
        except Exception as e:
            print(f"    ERROR: {e}")

    print()
    print(f"Done. {len(generators)} challenge files created in {OUTPUT_DIR}/")
    print()
    print("Existing challenge files (not regenerated):")
    existing = ["challenge_footnotes.docx", "challenge_nested_lists.docx",
                "challenge_speaker_notes.pptx", "challenge_formulas.xlsx",
                "challenge_complex.html"]
    for f in existing:
        path = OUTPUT_DIR / f
        if path.exists():
            print(f"  {path} ({path.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
